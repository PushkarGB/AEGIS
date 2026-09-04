"""Focused tests proving the real AEGIS document-drafting workflow through the runtime.

Validates the full execution path:
Request
→ RouterAgentRuntime
→ intent/modality/workflow decision (model determines workflow, no keyword matching)
→ ExecutionController
→ CapabilityBroker
→ extract_document (deterministic text extraction)
→ draft_approval_note (agent-role model drafting)
→ generate_word (DOCX generation)
→ verify_result (deterministic verification of structure & separation)
→ HITL approval (Controller-owned state machine)
→ FINAL DOCX

Invariants verified:
- Agent/model determines required workflow without keyword matching.
- Deterministic capabilities remain deterministic.
- Model drafting keeps extracted facts separate from recommendations.
- Approval remains an explicit controller/HITL state transition.
- Zero chain-of-thought exposed.
- Execution events and session/task/user identity preserved.
"""

from __future__ import annotations

import json
from pathlib import Path
from uuid import uuid4

import docx
import pytest

from aegis.agent import (
    AgentIntent,
    AttachmentDescriptor,
    InputModality,
    IntentAnalysisRequest,
    IntentAnalysisResult,
    ObservationDecision,
    RouterAgentRuntime,
)
from aegis.broker import RegistryCapabilityBroker
from aegis.capabilities import (
    CapabilityRegistry,
    DraftApprovalNoteCapability,
    ExtractDocumentCapability,
    FinishCapability,
    GenerateWordCapability,
    VerifyResultCapability,
)
from aegis.config import load_config
from aegis.events import (
    ExecutionEvent,
    ExecutionEventContext,
    ExecutionEventPublisher,
    ExecutionEventStatus,
    ExecutionEventType,
)
from aegis.orchestration import ExecutionController, RuntimeTaskRunner, WorkflowName
from aegis.orchestration.hitl import HITLApprovalState
from aegis.router import MockModelProvider, ModelRegistry, ModelRouter
from aegis.schemas import ApprovalStatus, FinalStatus, TaskState, VerificationStatus
from demo.fixtures import (
    EXPECTED_DRAFT_APPROVAL_NOTE,
    EXPECTED_INSPECTION_TEXT_SNIPPETS,
    SYNTHETIC_INSPECTION_REPORT_PDF,
)


USER_PROMPT = (
    "Review this inspection report and draft an approval note summarizing the key findings, "
    "supporting observations, and recommended actions. Keep extracted facts separate from "
    "recommendations and require operator approval before finalizing the note."
)


@pytest.fixture
def fixture_pdf_path() -> Path:
    assert SYNTHETIC_INSPECTION_REPORT_PDF.exists(), "Synthetic inspection report fixture must exist"
    return SYNTHETIC_INSPECTION_REPORT_PDF


@pytest.fixture
def mock_agent_provider() -> MockModelProvider:
    """Mock agent model provider that returns structured JSON decisions for intent and drafting."""
    agent_call = 0

    def agent_response(request) -> str:
        nonlocal agent_call
        agent_call += 1

        # Check prompt content to distinguish intent classification vs drafting
        prompt_text = getattr(request, "prompt", "")

        if "Classify the request" in prompt_text or agent_call == 1:
            return json.dumps(
                {
                    "intent": "document_drafting",
                    "modality": "scanned_document",
                    "workflow": "scanned_document_approval",
                    "summary": "The user requested an approval note review of an inspection report.",
                }
            )

        # Drafting request
        return json.dumps(EXPECTED_DRAFT_APPROVAL_NOTE)

    return MockModelProvider(response_factory=agent_response)


class TestDocumentDraftingWorkflow:
    """End-to-end tests for the document-drafting approval note workflow."""

    def test_full_document_drafting_workflow_execution(
        self,
        fixture_pdf_path: Path,
        mock_agent_provider: MockModelProvider,
        tmp_path: Path,
    ):
        """Prove the complete execution path from natural language request to verified DOCX with HITL."""
        config = load_config()
        publisher = ExecutionEventPublisher()
        router = ModelRouter(ModelRegistry(config.models))

        providers = {
            "local_ollama": mock_agent_provider,
        }

        agent_runtime = RouterAgentRuntime(
            config=config.agent,
            router=router,
            providers=providers,
            event_publisher=publisher,
        )

        deliverables_dir = tmp_path / "deliverables"

        registry = CapabilityRegistry(config.capabilities)
        registry.register(ExtractDocumentCapability())
        registry.register(
            DraftApprovalNoteCapability(router=router, providers=providers)
        )
        registry.register(GenerateWordCapability(output_dir=deliverables_dir))
        registry.register(VerifyResultCapability())
        registry.register(FinishCapability())

        runner = RuntimeTaskRunner(
            event_publisher=publisher,
            agent_runtime=agent_runtime,
            router=router,
            providers=providers,
            capability_registry=registry,
            deliverables_dir=deliverables_dir,
            config=config,
        )

        session_id = uuid4()
        task_id = uuid4()
        user_id = "alice"

        # 1. Start execution — should proceed through drafting and pause at HITL approval gate
        run_res = runner.start_execution(
            session_id=session_id,
            task_id=task_id,
            user_id=user_id,
            user_goal=USER_PROMPT,
            attachment_path=str(fixture_pdf_path),
        )

        # Verify task is currently awaiting operator approval
        assert run_res.hitl_state == HITLApprovalState.WAITING_FOR_APPROVAL
        assert run_res.final_status == FinalStatus.NOT_FINAL
        assert run_res.workflow_id == WorkflowName.SCANNED_DOCUMENT_APPROVAL.value
        assert len(run_res.artifact_paths) == 1

        controller = runner.get_controller(task_id)
        assert controller is not None
        assert controller.state.approval_status == ApprovalStatus.PENDING
        assert controller.state.verification_status == VerificationStatus.PASSED

        # 2. Check deliverable DOCX content
        docx_path = Path(run_res.artifact_paths[0])
        assert docx_path.exists()
        doc = docx.Document(str(docx_path))
        para_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(cell.text for tbl in doc.tables for row in tbl.rows for cell in row.cells)
        full_text = para_text + "\n" + table_text
        assert "1. Extracted Facts and Key Findings" in full_text
        assert "2. Supporting Observations" in full_text
        assert "3. Recommended Actions" in full_text
        assert "PUMP-104B" in full_text
        assert "TANK-301A" in full_text
        assert "DRAFT" in full_text

        # 3. Verify user result text distinguishes sections cleanly
        assert "1. Extracted Facts & Key Findings" in run_res.result_text
        assert "2. Supporting Observations" in run_res.result_text
        assert "3. Recommended Actions" in run_res.result_text
        assert "WAITING FOR OPERATOR APPROVAL" in run_res.result_text
        assert "<think>" not in run_res.result_text
        assert "chain_of_thought" not in run_res.result_text

        # 4. Operator grants human approval
        final_res = runner.record_approval(
            task_id=task_id,
            user_id="operator_bob",
            approved=True,
        )

        # 5. Verify final status is COMPLETED and state machine reached APPROVED
        assert final_res.final_status == FinalStatus.COMPLETED
        assert final_res.hitl_state == HITLApprovalState.APPROVED or final_res.hitl_state == HITLApprovalState.FINAL
        assert "APPROVED and FINALIZED" in final_res.result_text
        assert controller.state.approval_status == ApprovalStatus.APPROVED

        # 6. Verify all execution events preserve identity and no chain of thought
        events = final_res.events
        assert len(events) > 0
        event_types = [ev.event_type for ev in events]
        assert ExecutionEventType.DOCUMENT_TYPE_IDENTIFIED in event_types
        assert ExecutionEventType.TASK_STARTED in event_types
        assert ExecutionEventType.WORKFLOW_SELECTED in event_types
        assert ExecutionEventType.CAPABILITY_STARTED in event_types
        assert ExecutionEventType.VERIFICATION_STARTED in event_types
        assert ExecutionEventType.VERIFICATION_COMPLETED in event_types
        assert ExecutionEventType.HITL_REQUIRED in event_types
        assert ExecutionEventType.APPROVAL_RECORDED in event_types
        assert ExecutionEventType.TASK_COMPLETED in event_types

        for ev in events:
            assert ev.session_id == session_id
            assert ev.task_id == task_id
            assert "<think>" not in ev.summary

    def test_document_drafting_operator_rejection(
        self,
        fixture_pdf_path: Path,
        mock_agent_provider: MockModelProvider,
        tmp_path: Path,
    ):
        """Prove that operator rejection halts the workflow and marks the task cancelled/rejected."""
        config = load_config()
        publisher = ExecutionEventPublisher()
        router = ModelRouter(ModelRegistry(config.models))
        providers = {"local_ollama": mock_agent_provider}

        agent_runtime = RouterAgentRuntime(config.agent, router, providers, publisher)

        registry = CapabilityRegistry(config.capabilities)
        registry.register(ExtractDocumentCapability())
        registry.register(DraftApprovalNoteCapability(router=router, providers=providers))
        registry.register(GenerateWordCapability(output_dir=tmp_path / "deliverables"))
        registry.register(VerifyResultCapability())
        registry.register(FinishCapability())

        runner = RuntimeTaskRunner(
            event_publisher=publisher,
            agent_runtime=agent_runtime,
            router=router,
            providers=providers,
            capability_registry=registry,
            deliverables_dir=tmp_path / "deliverables",
            config=config,
        )

        task_id = uuid4()
        run_res = runner.start_execution(
            session_id=uuid4(),
            task_id=task_id,
            user_id="alice",
            user_goal=USER_PROMPT,
            attachment_path=str(fixture_pdf_path),
        )
        assert run_res.hitl_state == HITLApprovalState.WAITING_FOR_APPROVAL

        # Operator rejects
        rejected_res = runner.record_approval(
            task_id=task_id,
            user_id="operator_bob",
            approved=False,
        )

        assert rejected_res.final_status == FinalStatus.CANCELLED
        assert rejected_res.hitl_state == HITLApprovalState.REJECTED
        assert "rejected by operator" in rejected_res.result_text
        assert "[Status: REJECTED]" in rejected_res.result_text

    def test_extract_document_capability_standalone(self, fixture_pdf_path: Path):
        """Verify standalone execution of ExtractDocumentCapability on the inspection report."""
        cap = ExtractDocumentCapability()
        from aegis.schemas import CapabilityRequest

        req = CapabilityRequest(
            capability_name="extract_document",
            inputs={"document_path": str(fixture_pdf_path)},
        )
        result = cap.execute(req)
        assert result.status.value == "succeeded"
        assert result.output["page_count"] == 2
        text = result.output["text"]
        for snippet in EXPECTED_INSPECTION_TEXT_SNIPPETS:
            assert snippet in text, f"Snippet '{snippet}' should be extracted from PDF"

    def test_draft_approval_note_emits_model_invoked_event(self):
        """Verify DraftApprovalNoteCapability emits MODEL_INVOKED with prompt and raw response."""
        from aegis.events import ExecutionEventPublisher, ExecutionEventType
        from aegis.router import MockModelProvider, ModelRegistry, ModelRouter
        from aegis.schemas import CapabilityRequest

        config = load_config()
        publisher = ExecutionEventPublisher()
        provider = MockModelProvider(response_factory=lambda req: json.dumps(EXPECTED_DRAFT_APPROVAL_NOTE))
        router = ModelRouter(ModelRegistry(config.models))
        cap = DraftApprovalNoteCapability(
            router=router,
            providers={"local_ollama": provider},
            event_publisher=publisher,
        )

        req = CapabilityRequest(
            capability_name="draft_approval_note",
            inputs={
                "extracted_text": "Sample inspection text",
                "user_goal": "Draft note",
                "document_title": "Report",
            },
        )
        res = cap.execute(req)
        assert res.status.value == "succeeded"

        events = [e for e in publisher.events if e.event_type == ExecutionEventType.MODEL_INVOKED]
        assert len(events) == 1
        ev = events[0]
        assert ev.metadata["role"] == "agent"
        assert ev.metadata["task_type"] == "drafting"
        assert "Sample inspection text" in ev.metadata["prompt"]
        assert ev.metadata["model_raw_response"] is not None
