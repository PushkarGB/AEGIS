"""GenerateWord capability: produces formatted DOCX approval notes.

Uses python-docx to generate an executive approval note document with:
- Standardized header / metadata banner
- Clearly distinguished Extracted Facts/Findings section
- Supporting Observations section
- Recommended Actions section
- Human-in-the-Loop Approval Status stamp

Operates deterministically without cloud or external services.
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from uuid import uuid4

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from aegis.capabilities.base import (
    Capability,
    CapabilityContract,
    CapabilityKind,
    CapabilityMetadata,
)
from aegis.schemas import (
    Artifact,
    CapabilityRequest,
    CapabilityResult,
    CapabilityResultStatus,
    Observation,
)


def create_approval_note_docx(
    title: str,
    document_reference: str,
    key_findings: list[str],
    supporting_observations: list[str],
    recommendations: list[str],
    approval_status: str = "DRAFT — PENDING OPERATOR APPROVAL",
    summary: str | None = None,
    output_path: Path | str | None = None,
) -> Path:
    """Deterministically generate a formatted .docx approval note.

    Parameters
    ----------
    title:
        Document title.
    document_reference:
        Source document name/reference.
    key_findings:
        Extracted factual findings from inspection.
    supporting_observations:
        Supporting observations from review.
    recommendations:
        Actionable recommendations.
    approval_status:
        Approval status stamp (e.g. DRAFT or FINAL).
    summary:
        Executive summary text.
    output_path:
        Destination file path.

    Returns
    -------
    Path:
        Path to the saved DOCX file.
    """
    path = Path(output_path or f"deliverables/Approval_Note_{uuid4().hex[:8]}.docx")
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = docx.Document()

    # 1. Document Title
    h1 = doc.add_heading(level=0)
    title_run = h1.add_run(title)
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # 2. Metadata / Header Banner Table
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    hdr_data = [
        ("Source Document:", document_reference),
        ("Generated Date:", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")),
        ("Approval Status:", approval_status),
        ("Governance:", "AEGIS Autonomous Industrial Agent Core"),
    ]
    for idx, (lbl, val) in enumerate(hdr_data):
        row = table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        p_lbl = cell_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(lbl)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(9.5)
        p_val = cell_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9.5)
        if lbl == "Approval Status:":
            r_val.font.bold = True
            if "APPROVED" in val or "FINAL" in val:
                r_val.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            else:
                r_val.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph("")

    # 3. Executive Summary
    if summary:
        doc.add_heading("Executive Summary", level=2)
        p = doc.add_paragraph(summary)
        p.paragraph_format.line_spacing = 1.15

    # 4. Section: Extracted Facts & Key Findings (Clearly Separated)
    doc.add_heading("1. Extracted Facts and Key Findings", level=2)
    p_intro1 = doc.add_paragraph(
        "The following factual determinations were extracted directly from the inspection report:"
    )
    p_intro1.paragraph_format.italic = True
    for item in key_findings:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.size = Pt(10.5)

    # 5. Section: Supporting Observations
    doc.add_heading("2. Supporting Observations", level=2)
    p_intro2 = doc.add_paragraph(
        "Technical context and condition assessments observed during review:"
    )
    p_intro2.paragraph_format.italic = True
    for item in supporting_observations:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.size = Pt(10.5)

    # 6. Section: Recommended Actions
    doc.add_heading("3. Recommended Actions", level=2)
    p_intro3 = doc.add_paragraph(
        "Proposed maintenance actions requiring human operator review and authorization:"
    )
    p_intro3.paragraph_format.italic = True
    for item in recommendations:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.size = Pt(10.5)

    # 7. Sign-Off & Approval Block
    doc.add_heading("4. Operator Sign-Off & Approval", level=2)
    sign_table = doc.add_table(rows=3, cols=2)
    sign_table.style = "Table Grid"
    sign_data = [
        ("Reviewing Authority:", "Operations & Integrity Assurance"),
        ("Current Governance State:", approval_status),
        ("Operator Signature / Timestamp:", "Pending HITL Controller Confirmation" if "PENDING" in approval_status else "CONFIRMED & AUTHORIZED"),
    ]
    for idx, (lbl, val) in enumerate(sign_data):
        row = sign_table.rows[idx]
        row.cells[0].paragraphs[0].add_run(lbl).font.bold = True
        row.cells[1].paragraphs[0].add_run(val)

    doc.save(str(path))
    return path


class GenerateWordCapability(Capability):
    """Generate DOCX approval notes deterministically using python-docx."""

    def __init__(self, output_dir: Path | str | None = None) -> None:
        self._output_dir = Path(output_dir or "deliverables")
        self._metadata = CapabilityMetadata(
            name="generate_word",
            kind=CapabilityKind.TOOL,
            description="Generate a formatted DOCX approval note deliverable.",
            input_contract=CapabilityContract(
                json_schema={
                    "type": "object",
                    "properties": {
                        "title": {"type": "string"},
                        "document_reference": {"type": "string"},
                        "key_findings": {"type": "array", "items": {"type": "string"}},
                        "supporting_observations": {"type": "array", "items": {"type": "string"}},
                        "recommendations": {"type": "array", "items": {"type": "string"}},
                        "approval_status": {"type": "string"},
                        "summary": {"type": "string"},
                    },
                }
            ),
            output_contract=CapabilityContract(
                json_schema={
                    "type": "object",
                    "properties": {
                        "file_path": {"type": "string"},
                        "title": {"type": "string"},
                        "approval_status": {"type": "string"},
                    },
                    "required": ["file_path"],
                }
            ),
            input_modalities=("document", "scanned_document"),
        )

    @property
    def metadata(self) -> CapabilityMetadata:
        return self._metadata

    def execute(self, request: CapabilityRequest) -> CapabilityResult:
        inputs = request.inputs
        title = inputs.get("title", "APPROVAL NOTE: Equipment Inspection Review")
        doc_ref = inputs.get("document_reference", "Inspection Report")
        findings = inputs.get("key_findings", [])
        observations = inputs.get("supporting_observations", [])
        recommendations = inputs.get("recommendations", [])
        approval_status = inputs.get("approval_status", "DRAFT — PENDING OPERATOR APPROVAL")
        summary = inputs.get("summary", "")

        out_path = self._output_dir / f"Approval_Note_{uuid4().hex[:8]}.docx"
        try:
            saved_path = create_approval_note_docx(
                title=title,
                document_reference=doc_ref,
                key_findings=findings,
                supporting_observations=observations,
                recommendations=recommendations,
                approval_status=approval_status,
                summary=summary,
                output_path=out_path,
            )
        except Exception as exc:
            return CapabilityResult(
                request_id=request.request_id,
                status=CapabilityResultStatus.FAILED,
                error=f"DOCX artifact generation failed: {exc}",
            )

        artifact = Artifact(
            name=saved_path.name,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            location=str(saved_path),
            description="Draft approval note deliverable (.docx)",
            source_request_id=request.request_id,
        )

        obs = Observation(
            source="generate_word",
            kind="word_artifact_generated",
            summary=f"Generated Word document at {saved_path}.",
            data={
                "file_path": str(saved_path),
                "approval_status": approval_status,
            },
            request_id=request.request_id,
        )

        return CapabilityResult(
            request_id=request.request_id,
            status=CapabilityResultStatus.SUCCEEDED,
            output={
                "file_path": str(saved_path),
                "title": title,
                "approval_status": approval_status,
            },
            observations=[obs],
            artifacts=[artifact],
        )
